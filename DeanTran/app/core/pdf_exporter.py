"""
PDF Exporter – Apply page setup to Word/Excel/PPT documents, then export to PDF.

Supports:
- Word (.docx): margins, orientation, paper size, centering
- Excel (.xlsx): margins, orientation, paper size, centering, fit-to-page
- PowerPoint (.pptx): slide dimensions, orientation
- PDF export via Microsoft Office COM automation (pywin32)
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("DeanTran.pdf_exporter")

# Paper sizes in centimetres (width, height) – portrait orientation
PAPER_SIZES = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "Letter": (21.59, 27.94),
    "Legal": (21.59, 35.56),
    "B5": (17.6, 25.0),
}

# openpyxl paper-size enum values
_OPENPYXL_PAPER = {
    "Letter": "1",
    "Legal": "5",
    "A3": "8",
    "A4": "9",
    "A5": "11",
    "B5": "13",
}


def detect_auto_orientation(input_path: Path, ext: str) -> str:
    """Tự động phát hiện hướng trang tối ưu cho tài liệu (ngang/dọc)."""
    if ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(input_path), read_only=True)
            max_cols = 0
            for ws in wb.worksheets:
                if ws.max_column and ws.max_column > max_cols:
                    max_cols = ws.max_column
            wb.close()
            # Nếu có bất kỳ sheet nào nhiều hơn 8 cột, xoay ngang (landscape)
            if max_cols > 8:
                return "landscape"
        except Exception:
            pass
        return "portrait"
        
    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(input_path))
            # Nếu có bảng nào nhiều hơn 6 cột, xoay ngang
            for table in doc.tables:
                if len(table.columns) > 6:
                    return "landscape"
        except Exception:
            pass
        return "portrait"
        
    elif ext == ".pptx":
        return "landscape"
        
    return "portrait"


class PageSetup:
    """Container for page-setup parameters."""

    def __init__(
        self,
        paper_size: str = "A4",
        orientation: str = "portrait",       # portrait | landscape
        margin_top_cm: float = 2.54,
        margin_bottom_cm: float = 2.54,
        margin_left_cm: float = 3.18,
        margin_right_cm: float = 3.18,
        header_distance_cm: float = 1.27,
        footer_distance_cm: float = 1.27,
        center_horizontally: bool = False,
        center_vertically: bool = False,
        # Excel-only: fit to page
        fit_to_page: bool = False,
        fit_to_width: int = 1,     # number of pages wide (0 = auto)
        fit_to_height: int = 1,    # number of pages tall (0 = auto)
        # Excel-only: print options
        print_gridlines: bool = False,
        repeat_rows: int = 0,      # repeat first N rows on every page (0 = off)
    ):
        self.paper_size = paper_size
        self.orientation = orientation
        self.margin_top_cm = margin_top_cm
        self.margin_bottom_cm = margin_bottom_cm
        self.margin_left_cm = margin_left_cm
        self.margin_right_cm = margin_right_cm
        self.header_distance_cm = header_distance_cm
        self.footer_distance_cm = footer_distance_cm
        self.center_horizontally = center_horizontally
        self.center_vertically = center_vertically
        self.fit_to_page = fit_to_page
        self.fit_to_width = fit_to_width
        self.fit_to_height = fit_to_height
        self.print_gridlines = print_gridlines
        self.repeat_rows = repeat_rows


def _cm_to_inches(cm: float) -> float:
    return cm / 2.54


# ─── Word (.docx) ───────────────────────────────────────────────────────

def apply_page_setup_docx(doc_path: str | Path, setup: PageSetup, output_path: str | Path) -> Path:
    """Apply page setup to a Word document."""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm
    from lxml import etree

    doc_path = Path(doc_path)
    output_path = Path(output_path)
    doc = Document(str(doc_path))

    paper_w_cm, paper_h_cm = PAPER_SIZES.get(setup.paper_size, PAPER_SIZES["A4"])
    if setup.orientation == "landscape":
        paper_w_cm, paper_h_cm = paper_h_cm, paper_w_cm

    WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE if setup.orientation == "landscape" else WD_ORIENT.PORTRAIT
        section.page_width = Cm(paper_w_cm)
        section.page_height = Cm(paper_h_cm)
        section.top_margin = Cm(setup.margin_top_cm)
        section.bottom_margin = Cm(setup.margin_bottom_cm)
        section.left_margin = Cm(setup.margin_left_cm)
        section.right_margin = Cm(setup.margin_right_cm)
        section.header_distance = Cm(setup.header_distance_cm)
        section.footer_distance = Cm(setup.footer_distance_cm)

        # Vertical alignment
        if setup.center_vertically:
            sectPr = section._sectPr
            vAlign = sectPr.find(f"{{{WML_NS}}}vAlign")
            if vAlign is None:
                vAlign = etree.SubElement(sectPr, f"{{{WML_NS}}}vAlign")
            vAlign.set(f"{{{WML_NS}}}val", "center")

    doc.save(str(output_path))
    logger.info("Word page setup applied -> %s", output_path)
    return output_path


# ─── Excel (.xlsx) ──────────────────────────────────────────────────────

def apply_page_setup_excel(xlsx_path: str | Path, setup: PageSetup, output_path: str | Path) -> Path:
    """Apply page setup to an Excel workbook (all sheets)."""
    import openpyxl

    xlsx_path = Path(xlsx_path)
    output_path = Path(output_path)
    wb = openpyxl.load_workbook(str(xlsx_path))

    paper_code = _OPENPYXL_PAPER.get(setup.paper_size, "9")  # default A4

    for ws in wb.worksheets:
        # Paper size
        ws.page_setup.paperSize = paper_code

        # Orientation
        if setup.orientation == "landscape":
            ws.page_setup.orientation = "landscape"
        else:
            ws.page_setup.orientation = "portrait"

        # Margins (openpyxl uses inches)
        ws.page_margins.top = _cm_to_inches(setup.margin_top_cm)
        ws.page_margins.bottom = _cm_to_inches(setup.margin_bottom_cm)
        ws.page_margins.left = _cm_to_inches(setup.margin_left_cm)
        ws.page_margins.right = _cm_to_inches(setup.margin_right_cm)
        ws.page_margins.header = _cm_to_inches(setup.header_distance_cm)
        ws.page_margins.footer = _cm_to_inches(setup.footer_distance_cm)

        # Centering
        ws.print_options.horizontalCentered = setup.center_horizontally
        ws.print_options.verticalCentered = setup.center_vertically

        # Fit to page (Excel-specific)
        if setup.fit_to_page:
            ws.page_setup.fitToPage = True
            ws.page_setup.fitToWidth = setup.fit_to_width
            ws.page_setup.fitToHeight = setup.fit_to_height
            ws.page_setup.scale = None
        else:
            ws.page_setup.fitToPage = False

        # Gridlines
        ws.print_options.gridLines = setup.print_gridlines
        ws.print_options.gridLinesSet = True

        # Repeat rows (print title rows)
        if setup.repeat_rows > 0:
            ws.print_title_rows = f'1:{setup.repeat_rows}'

    wb.save(str(output_path))
    wb.close()
    logger.info("Excel page setup applied → %s", output_path)
    return output_path


# ─── PowerPoint (.pptx) ────────────────────────────────────────────────

def apply_page_setup_pptx(pptx_path: str | Path, setup: PageSetup, output_path: str | Path) -> Path:
    """Apply slide dimensions to a PowerPoint presentation."""
    from pptx import Presentation
    from pptx.util import Cm

    pptx_path = Path(pptx_path)
    output_path = Path(output_path)
    prs = Presentation(str(pptx_path))

    paper_w_cm, paper_h_cm = PAPER_SIZES.get(setup.paper_size, PAPER_SIZES["A4"])
    if setup.orientation == "landscape":
        paper_w_cm, paper_h_cm = paper_h_cm, paper_w_cm

    prs.slide_width = Cm(paper_w_cm)
    prs.slide_height = Cm(paper_h_cm)

    prs.save(str(output_path))
    logger.info("PPT slide size applied → %s", output_path)
    return output_path


# ─── PDF Conversion via COM ─────────────────────────────────────────────

def convert_to_pdf_com(source_path: str | Path, pdf_path: str | Path) -> Path:
    """
    Convert Word/Excel/PPT to PDF using Microsoft Office COM automation.
    Works for .docx, .xlsx, .pptx files.
    Requires Microsoft Office installed on Windows.
    """
    import pythoncom
    import win32com.client

    source_path = Path(source_path).resolve()
    pdf_path = Path(pdf_path).resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    ext = source_path.suffix.lower()

    pythoncom.CoInitialize()
    try:
        if ext == ".docx":
            _convert_word_to_pdf(source_path, pdf_path)
        elif ext == ".xlsx":
            _convert_excel_to_pdf(source_path, pdf_path)
        elif ext == ".pptx":
            _convert_ppt_to_pdf(source_path, pdf_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    finally:
        pythoncom.CoUninitialize()

    logger.info("PDF conversion done: %s → %s", source_path.name, pdf_path.name)
    return pdf_path


def _convert_word_to_pdf(docx_path: Path, pdf_path: Path):
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    try:
        # Open in normal write mode to modify wrap format if needed, then save as PDF
        doc = word.Documents.Open(str(docx_path))
        
        # wdWrapSquare = 0
        # Check all floating shapes and set wrap format to Square if they are overlapping (WrapFormat.Type == 3 or others)
        for shape in doc.Shapes:
            try:
                # If wrap format is currently "None" (behind/in front of text) which causes overlaps
                if shape.WrapFormat.Type == 3:  # 3 = wdWrapNone
                    shape.WrapFormat.Type = 0  # 0 = wdWrapSquare
            except Exception:
                pass

        # wdFormatPDF = 17
        doc.SaveAs2(str(pdf_path), FileFormat=17)
        doc.Close(True)  # Save changes to docx as well so format is fixed in Word doc
    except Exception as exc:
        logger.warning(f"Error during Word to PDF conversion: {exc}")
        raise exc
    finally:
        try:
            word.Quit(False)
        except Exception:
            pass


def _convert_excel_to_pdf(xlsx_path: Path, pdf_path: Path):
    import win32com.client
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    temp_xlsx_path = None
    try:
        wb = None
        try:
            # Step 1: Open with UpdateLinks=0
            wb = excel.Workbooks.Open(str(xlsx_path), UpdateLinks=0, ReadOnly=True)
        except Exception as exc1:
            logger.warning(f"Excel standard open failed for {xlsx_path.name}, trying CorruptLoad=2: {exc1}")
            try:
                # Step 2: Open with CorruptLoad=2 (xlRepairFile)
                wb = excel.Workbooks.Open(str(xlsx_path), UpdateLinks=0, ReadOnly=True, CorruptLoad=2)
            except Exception as exc2:
                logger.warning(f"Excel repair open failed for {xlsx_path.name}, trying openpyxl sanitization: {exc2}")
                try:
                    # Step 3: openpyxl sanitization (clean sheet titles)
                    import openpyxl
                    temp_xlsx_path = xlsx_path.parent / f"~dt_tmp_{xlsx_path.stem}.xlsx"
                    op_wb = openpyxl.load_workbook(str(xlsx_path))
                    for ws in op_wb.worksheets:
                        title = ws.title.strip().strip("'")
                        for char in ['\\', '/', '?', '*', '[', ']', ':']:
                            title = title.replace(char, ' ')
                        title = title.strip().strip("'").strip()
                        title = title[:31].strip().strip("'").strip()
                        ws.title = title if title else "Sheet"
                    op_wb.save(str(temp_xlsx_path))
                    op_wb.close()

                    # Open the sanitized temp workbook
                    wb = excel.Workbooks.Open(str(temp_xlsx_path), UpdateLinks=0, ReadOnly=True)
                except Exception as exc3:
                    logger.warning(f"Excel sanitization open failed for {xlsx_path.name}, trying data_only fallback: {exc3}")
                    try:
                        # Step 4: openpyxl sanitization with data_only=True (strips complex formulas/links)
                        import openpyxl
                        if temp_xlsx_path and temp_xlsx_path.exists():
                            try:
                                temp_xlsx_path.unlink()
                            except Exception:
                                pass
                        temp_xlsx_path = xlsx_path.parent / f"~dt_tmp_{xlsx_path.stem}.xlsx"
                        op_wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
                        for ws in op_wb.worksheets:
                            title = ws.title.strip().strip("'")
                            for char in ['\\', '/', '?', '*', '[', ']', ':']:
                                title = title.replace(char, ' ')
                            title = title.strip().strip("'").strip()
                            title = title[:31].strip().strip("'").strip()
                            ws.title = title if title else "Sheet"
                        op_wb.save(str(temp_xlsx_path))
                        op_wb.close()

                        # Open the sanitized data_only temp workbook
                        wb = excel.Workbooks.Open(str(temp_xlsx_path), UpdateLinks=0, ReadOnly=True)
                    except Exception as exc4:
                        logger.error(f"Excel data_only fallback open failed for {xlsx_path.name}: {exc4}")
                        raise exc4

        # xlTypePDF = 0
        try:
            wb.ExportAsFixedFormat(0, str(pdf_path))
        except Exception:
            wb.SaveAs(str(pdf_path), FileFormat=57)
        wb.Close(False)
    finally:
        try:
            excel.Quit()
        except Exception:
            pass
        if temp_xlsx_path and temp_xlsx_path.exists():
            try:
                temp_xlsx_path.unlink()
            except Exception:
                pass


def _convert_ppt_to_pdf(pptx_path: Path, pdf_path: Path):
    import win32com.client
    ppt = win32com.client.Dispatch("PowerPoint.Application")
    try:
        # Try WithWindow=False first, fallback to standard Open if WPS complains
        try:
            presentation = ppt.Presentations.Open(str(pptx_path), WithWindow=False)
        except Exception:
            presentation = ppt.Presentations.Open(str(pptx_path))
        # ppSaveAsPDF = 32
        presentation.SaveAs(str(pdf_path), FileFormat=32)
        presentation.Close()
    finally:
        try:
            ppt.Quit()
        except Exception:
            pass


# ─── Main Export Pipeline ────────────────────────────────────────────────

def _apply_setup(ext: str, input_path: Path, setup: PageSetup, output_path: Path):
    """Apply the appropriate page setup function based on file extension."""
    if ext == ".docx":
        apply_page_setup_docx(input_path, setup, output_path)
    elif ext == ".xlsx":
        apply_page_setup_excel(input_path, setup, output_path)
    elif ext == ".pptx":
        apply_page_setup_pptx(input_path, setup, output_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _log_setup(ext: str, setup: PageSetup, _log, is_original: bool = False):
    """Log the applied or original page setup details."""
    prefix = "Original page setup" if is_original else "Page setup applied"
    _log("INFO", f"{prefix}: paper={setup.paper_size}, "
         f"orientation={setup.orientation}, "
         f"margins T={setup.margin_top_cm} B={setup.margin_bottom_cm} "
         f"L={setup.margin_left_cm} R={setup.margin_right_cm}")
    if ext == ".xlsx" and setup.fit_to_page:
        _log("INFO", f"Scaling: width={setup.fit_to_width}, height={setup.fit_to_height}")

def export_document(
    input_path: str | Path,
    setup: PageSetup,
    output_format: str = "pdf",   # "pdf" | "apply"
    overwrite: bool = False,      # True = overwrite original file (apply mode only)
    output_dir: Optional[str | Path] = None,
    log_fn=None,
    keep_original: bool = False,
) -> Path:
    """
    Full export pipeline:
      - output_format == "pdf":  apply page setup → convert to PDF
      - output_format == "apply": apply page setup directly to the file
        - overwrite=False: saves as _formatted copy (safe)
        - overwrite=True: overwrites the original file

    Supports .docx, .xlsx, .pptx files.
    Returns the final output file path.
    """
    import shutil
    import tempfile

    input_path = Path(input_path)
    _log = log_fn or (lambda lvl, msg: logger.info("[%s] %s", lvl, msg))

    if output_dir is None:
        output_dir = input_path.parent
    else:
        output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    ext = input_path.suffix.lower()
    stem = input_path.stem

    # If keep_original is requested, read setup from the original file
    if keep_original:
        from app.core.page_setup_reader import read_page_setup
        setup = read_page_setup(input_path)
        _log("INFO", f"Chế độ giữ nguyên gốc được kích hoạt cho: {input_path.name}")
    else:
        # Auto orientation resolution
        if getattr(setup, 'orientation', 'portrait') == "auto":
            detected_orient = detect_auto_orientation(input_path, ext)
            from copy import copy
            setup = copy(setup)
            setup.orientation = detected_orient
            _log("INFO", f"Auto-detected orientation: {detected_orient.capitalize()}")

    # Determine output filename
    if output_format == "pdf":
        final_name = f"{stem}.pdf"
    elif overwrite:
        # Overwrite: write to a temp file first, then replace original
        final_name = None  # handled specially below
    else:
        # Safe copy
        final_name = f"{stem}_formatted{ext}"

    final_path = (input_path if final_name is None else output_dir / final_name)

    _log("INFO", f"Applying page setup to: {input_path.name}" if not keep_original else f"Processing (keeping original format): {input_path.name}")

    if output_format == "apply":
        # Apply settings to file
        if overwrite:
            # Overwrite original: write to temp first, then replace
            if keep_original:
                _log_setup(ext, setup, _log, is_original=True)
                _log("INFO", f"✅ File gốc giữ nguyên (không sửa đổi): {input_path.name}")
                final_path = input_path
            else:
                tmp_dir = Path(tempfile.mkdtemp(prefix="deantran_apply_"))
                tmp_file = tmp_dir / f"{stem}_tmp{ext}"
                try:
                    _apply_setup(ext, input_path, setup, tmp_file)
                    # Replace original with temp file
                    shutil.move(str(tmp_file), str(input_path))
                    final_path = input_path
                    _log_setup(ext, setup, _log)
                    _log("INFO", f"✅ File gốc đã được ghi đè: {input_path.name}")
                except Exception as exc:
                    _log("ERROR", f"Failed to apply settings: {exc}")
                    raise
                finally:
                    try:
                        shutil.rmtree(str(tmp_dir), ignore_errors=True)
                    except Exception:
                        pass
        else:
            # Safe copy: write to _formatted file
            try:
                if keep_original:
                    shutil.copy(str(input_path), str(final_path))
                    _log_setup(ext, setup, _log, is_original=True)
                    _log("INFO", f"✅ Bản sao đã lưu (giữ nguyên gốc): {final_path.name}")
                else:
                    _apply_setup(ext, input_path, setup, final_path)
                    _log_setup(ext, setup, _log)
                    _log("INFO", f"✅ Bản sao đã lưu: {final_path.name}")
            except Exception as exc:
                _log("ERROR", f"Failed to apply settings: {exc}")
                raise
    else:
        # PDF mode: apply settings to temp → convert to PDF
        tmp_dir = Path(tempfile.mkdtemp(prefix="deantran_export_"))
        tmp_file = tmp_dir / f"{stem}_setup{ext}"

        try:
            if keep_original:
                shutil.copy(str(input_path), str(tmp_file))
                _log_setup(ext, setup, _log, is_original=True)
            else:
                _apply_setup(ext, input_path, setup, tmp_file)
                _log_setup(ext, setup, _log)

            _log("INFO", "Converting to PDF via Microsoft Office…")
            try:
                convert_to_pdf_com(tmp_file, final_path)
                _log("INFO", f"✅ PDF exported: {final_path.name}")
            except Exception as exc:
                _log("ERROR", f"PDF conversion failed: {exc}")
                raise
        finally:
            try:
                shutil.rmtree(str(tmp_dir), ignore_errors=True)
            except Exception:
                pass

    return final_path


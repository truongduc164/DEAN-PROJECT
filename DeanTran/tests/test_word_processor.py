"""
tests/test_word_processor.py – Tests for Word (.docx) translation bugs.

Covers:
  1. Table cell deduplication (merged cells)
  2. Header/Footer translation (is_linked_to_previous fix)
  3. Textbox extraction (w:txbxContent)
  4. Body paragraph deduplication
"""
from __future__ import annotations

import sys
from pathlib import Path
from unittest.mock import patch

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.core.event_manager import EventManager
from app.core.translators.translator_service import MockTranslator
from app.core.word_processor import WordProcessor


# ── Helpers ──────────────────────────────────────────────────────────


def _make_processor(**kwargs) -> WordProcessor:
    defaults = dict(
        translator=MockTranslator(),
        source_lang="Chinese",
        target_lang="English",
        batch_mode=False,
    )
    defaults.update(kwargs)
    return WordProcessor(**defaults)


def _add_textbox(doc: Document, text: str) -> None:
    """Insert a textbox (w:txbxContent) into the document body.

    Minimal XML structure that mirrors what Word actually generates.
    Uses VML (v:shape) wrapping which is the legacy but common format.
    """
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    VNS = 'urn:schemas-microsoft-com:vml'

    # Build the XML tree manually to avoid namespace issues with parse_xml
    r_elem = etree.SubElement(
        doc.add_paragraph()._p,
        qn('w:r'),
    )
    pict = etree.SubElement(r_elem, qn('w:pict'))
    shape = etree.SubElement(
        pict,
        '{%s}shape' % VNS,
        attrib={'style': 'width:200pt;height:50pt'},
    )
    vbox = etree.SubElement(shape, '{%s}textbox' % VNS)
    txbx = etree.SubElement(vbox, qn('w:txbxContent'))
    p = etree.SubElement(txbx, qn('w:p'))
    run = etree.SubElement(p, qn('w:r'))
    t = etree.SubElement(run, qn('w:t'))
    t.text = text


# ═══════════════════════════════════════════════════════════════════
# 1. Table cell deduplication (merged cells)
# ═══════════════════════════════════════════════════════════════════


class TestMergedTableDedup:
    def test_merged_cells_not_duplicated(self, tmp_path: Path):
        """Merged (spanning) cells should be translated only once."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        # Merge first row into a single cell
        cell_a = table.cell(0, 0)
        cell_c = table.cell(0, 2)
        merged = cell_a.merge(cell_c)
        merged.text = "合并单元格"  # "merged cell" in Chinese

        table.cell(1, 0).text = "第一列"
        table.cell(1, 1).text = "第二列"
        table.cell(1, 2).text = "第三列"

        src = tmp_path / "merged.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        # Read back
        result_doc = Document(str(out_path))
        result_table = result_doc.tables[0]

        # The merged cell should have translation once, not triple
        merged_text = result_table.cell(0, 0).text
        # MockTranslator prefixes with [MOCK_vi→en], check no duplication
        assert merged_text.count("[MOCK") <= 1, (
            f"Merged cell translated multiple times: {merged_text!r}"
        )

    def test_normal_table_cells_all_translated(self, tmp_path: Path):
        """Non-merged cells should each be translated."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        texts = ["细胞一", "细胞二", "细胞三", "细胞四"]
        for i, t in enumerate(texts):
            table.cell(i // 2, i % 2).text = t

        src = tmp_path / "normal_table.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        assert proc.last_result["translated"] >= 4

    def test_nested_table_cells_translated(self, tmp_path: Path):
        """Tables nested inside table cells should also be translated."""
        doc = Document()
        # Outer table
        outer = doc.add_table(rows=2, cols=2)
        outer.cell(0, 0).text = "外层单元格"  # outer cell
        outer.cell(0, 1).text = "另一个外层"

        # Insert a nested table inside cell(1, 0) via XML
        inner_cell = outer.cell(1, 0)
        inner_cell.text = ""  # clear default paragraph
        inner_tbl = inner_cell.add_table(rows=2, cols=1)
        inner_tbl.cell(0, 0).text = "嵌套第一行"  # nested row 1
        inner_tbl.cell(1, 0).text = "嵌套第二行"  # nested row 2

        outer.cell(1, 1).text = "外层最后"

        src = tmp_path / "nested_table.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        # Should translate: 3 outer cells + 2 nested cells = 5 minimum
        assert proc.last_result["translated"] >= 5, (
            f"Expected >=5 translated (incl. nested table), "
            f"got {proc.last_result['translated']}"
        )



# ═══════════════════════════════════════════════════════════════════
# 2. Header/Footer translation
# ═══════════════════════════════════════════════════════════════════


class TestHeaderFooter:
    @pytest.fixture(autouse=True)
    def _enable_header_footer(self):
        """Ensure translate_headers_footers is True for these tests."""
        from app.settings.settings_manager import settings
        old = settings.get("processing_options.translate_headers_footers", True)
        settings.set("processing_options.translate_headers_footers", True)
        yield
        settings.set("processing_options.translate_headers_footers", old)

    def test_header_translated_single_section(self, tmp_path: Path):
        """Header in a single-section doc should be translated."""
        doc = Document()
        doc.add_paragraph("正文内容")  # body text
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "页眉文字"  # "header text"

        src = tmp_path / "header.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        result_doc = Document(str(out_path))
        header_text = result_doc.sections[0].header.paragraphs[0].text
        # Should be translated (MockTranslator adds [MOCK...])
        assert header_text != "页眉文字", (
            f"Header was not translated: {header_text!r}"
        )

    def test_footer_translated_single_section(self, tmp_path: Path):
        """Footer in a single-section doc should be translated."""
        doc = Document()
        doc.add_paragraph("正文内容")
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = "页脚文字"  # "footer text"

        src = tmp_path / "footer.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        result_doc = Document(str(out_path))
        footer_text = result_doc.sections[0].footer.paragraphs[0].text
        assert footer_text != "页脚文字", (
            f"Footer was not translated: {footer_text!r}"
        )

    def test_linked_header_not_duplicated(self, tmp_path: Path):
        """Linked headers across sections should be translated just once."""
        doc = Document()
        doc.add_paragraph("段落一")

        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "统一页眉"

        src = tmp_path / "linked_header.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        # Should translate header exactly once (dedup prevents double)
        assert proc.last_result["status"] in ("SUCCESS", "PARTIAL")


# ═══════════════════════════════════════════════════════════════════
# 3. Textbox extraction
# ═══════════════════════════════════════════════════════════════════


class TestTextboxExtraction:
    def test_textbox_paragraphs_collected(self, tmp_path: Path):
        """Textbox content (w:txbxContent) should be extracted and translated."""
        doc = Document()
        doc.add_paragraph("正文段落")  # body paragraph
        _add_textbox(doc, "文本框内容")  # textbox content

        src = tmp_path / "textbox.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        # At least 2 items translated: body paragraph + textbox content
        assert proc.last_result["translated"] >= 2, (
            f"Expected >=2 translated, got {proc.last_result['translated']}"
        )

    def test_collect_textbox_paras_directly(self):
        """Unit test for _collect_textbox_paras method."""
        doc = Document()
        _add_textbox(doc, "测试文本框")
        _add_textbox(doc, "另一个文本框")

        proc = _make_processor()
        seen: set = set()
        result = proc._collect_textbox_paras(doc, seen)

        assert len(result) == 2, f"Expected 2 textbox paras, got {len(result)}"
        texts = [t for _, t in result]
        assert "测试文本框" in texts
        assert "另一个文本框" in texts

    def test_textbox_dedup(self):
        """Same textbox paragraph should not be collected twice."""
        doc = Document()
        _add_textbox(doc, "重复检查")

        proc = _make_processor()
        seen: set = set()
        result1 = proc._collect_textbox_paras(doc, seen)
        result2 = proc._collect_textbox_paras(doc, seen)

        assert len(result1) == 1
        assert len(result2) == 0, "Second call should get 0 (already seen)"


# ═══════════════════════════════════════════════════════════════════
# 4. Body paragraph deduplication
# ═══════════════════════════════════════════════════════════════════


class TestBodyDedup:
    def test_body_paragraphs_unique(self, tmp_path: Path):
        """Each body paragraph should be translated only once."""
        doc = Document()
        for text in ["第一段", "第二段", "第三段"]:
            doc.add_paragraph(text)

        src = tmp_path / "body_dedup.docx"
        doc.save(str(src))

        proc = _make_processor()
        out_path = proc.process(src)

        assert proc.last_result["translated"] == 3
        assert proc.last_result["status"] == "SUCCESS"

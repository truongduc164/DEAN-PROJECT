from __future__ import annotations

import sys
from pathlib import Path
import pytest

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm

import openpyxl
from pptx import Presentation

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from app.core.pdf_exporter import PageSetup, export_document
from app.core.page_setup_reader import read_page_setup


def test_docx_page_setup_reader(tmp_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4 Landscape width
    section.page_height = Cm(21.0)  # A4 Landscape height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    file_path = tmp_path / "test_doc.docx"
    doc.save(str(file_path))

    setup = read_page_setup(file_path)
    assert setup.paper_size == "A4"
    assert setup.orientation == "landscape"
    assert setup.margin_top_cm == 1.5
    assert setup.margin_bottom_cm == 1.5
    assert setup.margin_left_cm == 2.0
    assert setup.margin_right_cm == 2.0


def test_xlsx_page_setup_reader(tmp_path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.page_setup.paperSize = "9"  # A4
    ws.page_setup.orientation = "landscape"
    ws.page_margins.top = 0.59  # 1.5 cm in inches (approx)
    ws.page_margins.bottom = 0.59
    ws.page_margins.left = 0.787  # 2.0 cm in inches
    ws.page_margins.right = 0.787
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 2
    ws.print_options.gridLines = True

    file_path = tmp_path / "test_sheet.xlsx"
    wb.save(str(file_path))

    setup = read_page_setup(file_path)
    assert setup.paper_size == "A4"
    assert setup.orientation == "landscape"
    assert abs(setup.margin_top_cm - 1.5) < 0.1
    assert abs(setup.margin_left_cm - 2.0) < 0.1
    assert setup.fit_to_page is True
    assert setup.fit_to_width == 1
    assert setup.fit_to_height == 2
    assert setup.print_gridlines is True


def test_pptx_page_setup_reader(tmp_path: Path):
    prs = Presentation()
    # A4: 21.0 x 29.7 cm
    # 29.7 cm = 10,692,000 EMU, 21.0 cm = 7,560,000 EMU
    prs.slide_width = 10692000
    prs.slide_height = 7560000

    file_path = tmp_path / "test_pres.pptx"
    prs.save(str(file_path))

    setup = read_page_setup(file_path)
    assert setup.paper_size == "A4"
    assert setup.orientation == "landscape"


def test_export_document_keep_original(tmp_path: Path):
    # Test Word document keep_original format (safe copy)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4 Portrait
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)

    file_path = tmp_path / "input.docx"
    doc.save(str(file_path))

    # We export with keep_original=True and output_format="apply" without overwrite
    # It should copy the file as _formatted.docx and keep the original settings.
    out_path = export_document(
        input_path=file_path,
        setup=PageSetup(paper_size="A3"),  # user requested A3 but we keep original!
        output_format="apply",
        overwrite=False,
        output_dir=tmp_path,
        keep_original=True,
    )

    assert out_path.exists()
    assert out_path != file_path
    assert out_path.name == "input_formatted.docx"

    # Read back formatted document to check if it kept the original page size (A4, not A3)
    setup = read_page_setup(out_path)
    assert setup.paper_size == "A4"
    assert setup.margin_top_cm == 2.0


def test_export_custom_output_dir(tmp_path: Path):
    doc = Document()
    file_path = tmp_path / "input.docx"
    doc.save(str(file_path))

    # We specify a different output directory
    custom_dir = tmp_path / "custom_output"
    custom_dir.mkdir()

    out_path = export_document(
        input_path=file_path,
        setup=PageSetup(paper_size="A4"),
        output_format="pdf",
        overwrite=False,
        output_dir=custom_dir,
    )

    assert out_path.exists()
    assert out_path.parent == custom_dir
    assert out_path.suffix == ".pdf"
